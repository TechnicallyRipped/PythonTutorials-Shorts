



from docx import Document

doc = Document()

doc.add_heading('My Document Title', level=1)
doc.add_heading('My Document Title 2', level=2)
doc.add_heading('My Document Title 3', level=3)

doc.add_paragraph('This is a paragraph.')
doc.add_paragraph('This is paragraph 2.')
doc.add_paragraph("This is the longest paragraph you've ever seen. If you're reading this please like and subscribe to my channel! I'm trying to hit 1k by the end of the year.")

doc.save('example_document.docx')




